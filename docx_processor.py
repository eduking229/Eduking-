"""
DocShield – Insertion du code-barres Code 128 dans les documents DOCX.
"""
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_horizontal_rule(paragraph):
    """Ajoute une ligne horizontale via XML OOXML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def stamp_docx(input_bytes: bytes, barcode_png: bytes, metadata: dict, position: str = "bottom") -> bytes:
    """
    Insère le code-barres dans le document DOCX.
    position: 'top' | 'bottom'
    Retourne les bytes du DOCX modifié.
    """
    doc = Document(io.BytesIO(input_bytes))

    bc_stream = io.BytesIO(barcode_png)
    bc_width = Inches(5.5)  # ~14 cm, centré

    if position == "top":
        # Insérer au début
        _insert_barcode_block_top(doc, bc_stream, bc_width, metadata)
    else:
        # Insérer à la fin
        _insert_barcode_block_bottom(doc, bc_stream, bc_width, metadata)

    out_buf = io.BytesIO()
    doc.save(out_buf)
    out_buf.seek(0)
    return out_buf.read()


def _insert_barcode_block_bottom(doc: Document, bc_stream, bc_width, metadata: dict):
    """Ajoute le bloc code-barres à la fin du document."""
    # Séparateur
    sep = doc.add_paragraph()
    _add_horizontal_rule(sep)
    sep.paragraph_format.space_before = Pt(12)

    # Titre DocShield
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("🔒 Document Certifié DocShield")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Image code-barres
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bc_stream.seek(0)
    run_img = img_p.add_run()
    run_img.add_picture(bc_stream, width=bc_width)

    # Métadonnées
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        f"ID: {metadata['doc_id']}  •  Signé: {metadata['timestamp']}  •  SIG: {metadata['signature']}"
    )
    meta_run.font.size = Pt(7)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _insert_barcode_block_top(doc: Document, bc_stream, bc_width, metadata: dict):
    """Insère le bloc code-barres au début du document."""
    from docx.oxml import OxmlElement

    # On ajoute les paragraphes puis on les déplace en tête
    paragraphs_to_add = []

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("🔒 Document Certifié DocShield")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    paragraphs_to_add.append(title_p)

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bc_stream.seek(0)
    run_img = img_p.add_run()
    run_img.add_picture(bc_stream, width=bc_width)
    paragraphs_to_add.append(img_p)

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        f"ID: {metadata['doc_id']}  •  Signé: {metadata['timestamp']}  •  SIG: {metadata['signature']}"
    )
    meta_run.font.size = Pt(7)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    paragraphs_to_add.append(meta_p)

    sep = doc.add_paragraph()
    _add_horizontal_rule(sep)
    sep.paragraph_format.space_after = Pt(12)
    paragraphs_to_add.append(sep)

    # Déplacer au début du body
    body = doc.element.body
    first_elem = body[0]
    for p in reversed(paragraphs_to_add):
        body.remove(p._element)
        body.insert(0, p._element)
